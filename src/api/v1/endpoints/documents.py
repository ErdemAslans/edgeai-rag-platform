"""Document management endpoints."""

import os
import uuid
from typing import Annotated, Any, Dict, List, Optional
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status, Query, BackgroundTasks
from sqlalchemy.ext.asyncio import AsyncSession
import structlog

from src.api.deps import get_db, get_embedding_service, CurrentUser, DbSession
from src.api.v1.schemas.documents import (
    DocumentResponse,
    DocumentListResponse,
    DocumentStatusResponse,
    ChunkResponse,
)
from src.config import settings
from src.core.di import get_container
from src.db.repositories.document import DocumentRepository
from src.db.repositories.chunk import ChunkRepository

logger = structlog.get_logger()

router = APIRouter()


def validate_file_extension(filename: str) -> bool:
    """Validate file extension against allowed extensions."""
    if not filename:
        return False
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext in settings.ALLOWED_EXTENSIONS


def validate_file_size(file_size: int) -> bool:
    """Validate file size against maximum allowed size."""
    max_size = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
    return file_size <= max_size


@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    background_tasks: BackgroundTasks,
    current_user: CurrentUser,
    db: Annotated[AsyncSession, Depends(get_db)],
    file: UploadFile = File(...),
) -> DocumentResponse:
    """Upload a new document for processing."""
    # Validate file extension
    if not validate_file_extension(file.filename):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type not allowed. Allowed types: {settings.ALLOWED_EXTENSIONS}",
        )
    
    # Read file content to check size
    content = await file.read()
    if not validate_file_size(len(content)):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File size exceeds maximum allowed size of {settings.MAX_UPLOAD_SIZE_MB}MB",
        )
    
    # Reset file position
    await file.seek(0)
    
    # Generate unique filename
    file_ext = file.filename.rsplit(".", 1)[-1].lower()
    unique_filename = f"{uuid.uuid4()}.{file_ext}"
    
    # Ensure upload directory exists
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)
    
    # Save file
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Create document record
    doc_repo = DocumentRepository(db)
    document = await doc_repo.create({
        "user_id": current_user.id,
        "filename": file.filename,
        "content_type": file.content_type or "application/octet-stream",
        "file_path": file_path,
        "file_size": len(content),
    })
    
    # Commit the document creation before background task
    await db.commit()
    
    # Trigger background document processing (don't pass db session - task creates its own)
    background_tasks.add_task(process_document_task, document.id)
    
    return DocumentResponse.model_validate(document)


async def process_document_task(document_id: uuid.UUID) -> None:
    """Background task to process a document.

    Creates its own database session to avoid issues with request-scoped sessions.
    Uses the DI container to get the embedding service singleton.
    """
    from src.db.session import async_session_factory

    async with async_session_factory() as db:
        try:
            doc_repo = DocumentRepository(db)
            chunk_repo = ChunkRepository(db)
            # Get embedding service from DI container (singleton)
            container = get_container()
            embedding_service = container.resolve("embedding_service")
            
            # Get document
            document = await doc_repo.get_by_id(document_id)
            if not document:
                logger.error("Document not found for processing", document_id=str(document_id))
                return
            
            # Update status to processing
            await doc_repo.update_status(document_id, "processing")
            await db.commit()
            
            # Extract text from document
            file_path = Path(document.file_path)
            if not file_path.exists():
                await doc_repo.update_status(document_id, "failed", "File not found on disk")
                await db.commit()
                return
            
            try:
                text_content = await extract_text_from_file(file_path)
            except Exception as e:
                await doc_repo.update_status(document_id, "failed", f"Text extraction failed: {str(e)}")
                await db.commit()
                return
            
            # Split text into chunks
            chunks = split_text_into_chunks(text_content)
            
            if not chunks:
                await doc_repo.update_status(document_id, "failed", "No content could be extracted")
                await db.commit()
                return
            
            # Delete existing chunks if reprocessing
            await chunk_repo.delete_by_document_id(document_id)
            
            # Generate embeddings and create chunks
            for i, chunk_text in enumerate(chunks):
                try:
                    embedding = await embedding_service.embed_text(chunk_text)
                    await chunk_repo.create({
                        "document_id": document_id,
                        "chunk_index": i,
                        "content": chunk_text,
                        "embedding": embedding,
                        "token_count": len(chunk_text.split()),
                        "chunk_metadata": {"source": document.filename, "page": i // 3 + 1},
                    })
                except Exception as e:
                    logger.error("Failed to process chunk", chunk_index=i, error=str(e))
                    continue
            
            # Update document status and chunk count
            chunk_count = await chunk_repo.count_by_document(document_id)
            await doc_repo.update(document_id, {"chunk_count": chunk_count, "status": "completed"})
            await db.commit()
            
            logger.info("Document processed successfully", document_id=str(document_id), chunk_count=chunk_count)
            
        except Exception as e:
            logger.error("Document processing failed", document_id=str(document_id), error=str(e))
            try:
                await db.rollback()
                await doc_repo.update_status(document_id, "failed", str(e))
                await db.commit()
            except Exception:
                pass


async def extract_text_from_file(file_path: Path) -> str:
    """Extract text content from a file.
    
    Strategy:
    1. For PDFs: Try fast pypdf first, use Docling OCR only if needed
    2. For DOCX/PPTX: Use Docling (fast, no OCR needed)
    3. For images: Use Docling with OCR
    """
    ext = file_path.suffix.lower()
    logger.info("Extracting text from file", file_path=str(file_path), extension=ext)
    
    # PDF files - try fast extraction first
    if ext == ".pdf":
        # First try pypdf (fast, for text-based PDFs)
        text = await extract_text_with_pypdf(file_path)
        if text and len(text.strip()) > 100:  # If we got meaningful text
            logger.info("Fast PDF extraction successful", content_length=len(text))
            return text
        
        # If pypdf failed or got little text, use Docling with OCR
        logger.info("PDF appears to be scanned, using Docling OCR")
        return await extract_with_docling(file_path, enable_ocr=True)
    
    # DOCX - Try Docling first, fallback to python-docx
    elif ext == ".docx":
        try:
            return await extract_with_docling(file_path, enable_ocr=False)
        except Exception as e:
            logger.warning("Docling failed for DOCX, trying python-docx", error=str(e))
            return await extract_docx_fallback(file_path)
    
    # PPTX - Use Docling
    elif ext == ".pptx":
        return await extract_with_docling(file_path, enable_ocr=False)
    
    # Images - Use Docling with OCR
    elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        return await extract_with_docling(file_path, enable_ocr=True)
    
    # Plain text files
    elif ext in {".txt", ".md"}:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    # JSON files
    elif ext == ".json":
        import json
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            return json.dumps(data, indent=2)
    
    # CSV files
    elif ext == ".csv":
        import csv
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            rows = [", ".join(row) for row in reader]
            return "\n".join(rows)
    
    # Excel files
    elif ext in {".xlsx", ".xls"}:
        try:
            import pandas as pd
            df = pd.read_excel(file_path)
            return df.to_string()
        except ImportError:
            raise Exception("pandas/openpyxl not installed for Excel processing")
    
    # HTML files
    elif ext in {".html", ".htm"}:
        try:
            from bs4 import BeautifulSoup
            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                return soup.get_text(separator="\n", strip=True)
        except ImportError:
            raise Exception("beautifulsoup4 not installed for HTML processing")
    
    else:
        raise Exception(f"Unsupported file type: {ext}")


async def extract_text_with_pypdf(file_path: Path) -> str:
    """Fallback PDF extraction using pypdf."""
    try:
        import pypdf
        text_parts = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            logger.info("pypdf opened PDF", page_count=len(reader.pages))
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.error("Failed to extract page", page=i, error=str(e))
                    continue
        
        full_text = "\n\n".join(text_parts)
        logger.info("pypdf extraction complete", total_length=len(full_text))
        return full_text
    except ImportError:
        raise Exception("pypdf library not installed")
    except Exception as e:
        raise Exception(f"pypdf extraction failed: {str(e)}")


async def extract_docx_fallback(file_path: Path) -> str:
    """Fallback DOCX extraction using python-docx."""
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(str(file_path))
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        content = "\n\n".join(paragraphs)
        logger.info("python-docx extraction complete", content_length=len(content))
        return content
        
    except ImportError:
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")


async def extract_with_docling(file_path: Path, enable_ocr: bool = False) -> str:
    """Extract text using IBM Docling for advanced document processing.
    
    Supports: PDF, DOCX, PPTX, HTML, images with OCR.
    """
    try:
        from docling.document_converter import DocumentConverter
        
        logger.info("Using Docling for document conversion", file_path=str(file_path), ocr_enabled=enable_ocr)
        
        # Create converter with GPU support if available
        try:
            from docling.datamodel.pipeline_options import PdfPipelineOptions, AcceleratorOptions
            from docling.datamodel.base_models import InputFormat
            from docling.document_converter import PdfFormatOption
            import torch

            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = enable_ocr
            pipeline_options.do_table_structure = True
            
            if torch.cuda.is_available():
                pipeline_options.accelerator_options = AcceleratorOptions(
                    num_threads=4, device="cuda"
                )
                logger.info("Enabled GPU acceleration for Docling", device="cuda")

            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
                }
            )
        except Exception as e:
            logger.warning("Could not configure advanced Docling options, using defaults", error=str(e))
            converter = DocumentConverter()
        
        # Convert document
        result = converter.convert(str(file_path))
        
        # Try markdown export first (preserves structure)
        try:
            markdown_content = result.document.export_to_markdown()
            if markdown_content and markdown_content.strip():
                logger.info(
                    "Docling extraction complete (markdown)",
                    file_path=str(file_path),
                    content_length=len(markdown_content),
                )
                return markdown_content
        except Exception as md_err:
            logger.warning("Markdown export failed, trying text", error=str(md_err))
        
        # Fallback to plain text
        try:
            text_content = result.document.export_to_text()
            if text_content and text_content.strip():
                logger.info(
                    "Docling extraction complete (text)",
                    file_path=str(file_path),
                    content_length=len(text_content),
                )
                return text_content
        except Exception as txt_err:
            logger.warning("Text export failed", error=str(txt_err))
        
        raise Exception("Docling could not extract any content from document")
        
    except ImportError as e:
        logger.error("Docling not installed", error=str(e))
        raise Exception(
            f"Docling library not available. Install with: pip install docling. Error: {str(e)}"
        )
    except Exception as e:
        error_msg = str(e)
        logger.error("Docling extraction failed", error=error_msg, file_path=str(file_path))
        
        # Provide helpful error messages
        if "torch" in error_msg.lower() or "cuda" in error_msg.lower():
            raise Exception(
                f"Docling GPU/PyTorch error. Try setting CUDA_VISIBLE_DEVICES='' for CPU mode. Error: {error_msg}"
            )
        elif "memory" in error_msg.lower():
            raise Exception(
                f"Out of memory during document processing. Try a smaller file. Error: {error_msg}"
            )
        else:
            raise Exception(f"Document parsing failed: {error_msg}")


def split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks."""
    if not text or not text.strip():
        return []
    
    # Clean text
    text = text.strip()
    
    # If text is short enough, return as single chunk
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at sentence boundary
        if end < len(text):
            # Look for sentence endings
            for sep in [". ", ".\n", "? ", "!\n", "\n\n"]:
                last_sep = text[start:end].rfind(sep)
                if last_sep != -1 and last_sep > chunk_size // 2:
                    end = start + last_sep + len(sep)
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start with overlap
        start = end - overlap if end < len(text) else end
    
    return chunks


@router.get("/", response_model=DocumentListResponse)
async def list_documents(
    current_user: CurrentUser,
    db: Annotated[AsyncSession, Depends(get_db)],
    skip: int = Query(0, ge=0),
    limit: int = Query(10, ge=1, le=100),
    status_filter: Optional[str] = Query(None, alias="status"),
) -> DocumentListResponse:
    """List all documents for the current user."""
    doc_repo = DocumentRepository(db)
    
    documents = await doc_repo.get_by_user_id(
        user_id=current_user.id,
        skip=skip,
        limit=limit,
    )
    total = await doc_repo.count_by_user(user_id=current_user.id)
    
    return DocumentListResponse(
        documents=[DocumentResponse.model_validate(doc) for doc in documents],
        total=total,
        skip=skip,
        limit=limit,
    )


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: uuid.UUID,
    current_user: CurrentUser,
    db: Annotated[AsyncSession, Depends(get_db)],
) -> DocumentResponse:
    """Get a specific document by ID."""
    doc_repo = DocumentRepository(db)
    document = await doc_repo.get_by_id(document_id)
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )
    
    if document.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not authorized to access this document",
        )
    
    return DocumentResponse.model_validate(document)


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: uuid.UUID,
    current_user: CurrentUser,
    db: Annotated[AsyncSession, Depends(get_db)],
) -> None:
    """Delete a document."""
    doc_repo = DocumentRepository(db)
    document = await doc_repo.get_by_id(document_id)
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )
    
    if document.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not authorized to delete this document",
        )
    
    # Delete file from storage
    if os.path.exists(document.file_path):
        os.remove(document.file_path)
    
    # Delete document record (cascades to chunks)
    await doc_repo.delete(document_id)


@router.get("/{document_id}/chunks", response_model=List[ChunkResponse])
async def get_document_chunks(
    document_id: uuid.UUID,
    current_user: CurrentUser,
    db: Annotated[AsyncSession, Depends(get_db)],
) -> List[ChunkResponse]:
    """Get all chunks for a document."""
    doc_repo = DocumentRepository(db)
    document = await doc_repo.get_by_id(document_id)
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )
    
    if document.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not authorized to access this document",
        )
    
    chunks = await doc_repo.get_chunks(document_id)
    return [ChunkResponse.model_validate(chunk) for chunk in chunks]


@router.post("/{document_id}/process", response_model=DocumentStatusResponse)
async def process_document(
    document_id: uuid.UUID,
    current_user: CurrentUser,
    background_tasks: BackgroundTasks,
    db: Annotated[AsyncSession, Depends(get_db)],
) -> DocumentStatusResponse:
    """Process a document (extract text, create chunks, generate embeddings)."""
    doc_repo = DocumentRepository(db)
    document = await doc_repo.get_by_id(document_id)
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )
    
    if document.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not authorized to process this document",
        )
    
    if document.status == "processing":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document is already being processed",
        )
    
    # Update status to pending
    await doc_repo.update_status(document_id, "pending")
    await db.commit()
    
    # Trigger background processing (don't pass db session - task creates its own)
    background_tasks.add_task(process_document_task, document_id)
    
    return DocumentStatusResponse(
        document_id=document_id,
        status="pending",
        message="Document queued for processing",
    )


@router.post("/{document_id}/reprocess", response_model=DocumentStatusResponse)
async def reprocess_document(
    document_id: uuid.UUID,
    current_user: CurrentUser,
    background_tasks: BackgroundTasks,
    db: Annotated[AsyncSession, Depends(get_db)],
) -> DocumentStatusResponse:
    """Reprocess a document (re-run ingestion pipeline)."""
    doc_repo = DocumentRepository(db)
    document = await doc_repo.get_by_id(document_id)
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )
    
    if document.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not authorized to reprocess this document",
        )
    
    if document.status == "processing":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document is already being processed",
        )
    
    # Update status to pending
    await doc_repo.update_status(document_id, "pending")
    await db.commit()
    
    # Trigger background processing (don't pass db session - task creates its own)
    background_tasks.add_task(process_document_task, document_id)
    
    return DocumentStatusResponse(
        document_id=document_id,
        status="pending",
        message="Document queued for reprocessing",
    )


@router.get("/{document_id}/status", response_model=DocumentStatusResponse)
async def get_document_status(
    document_id: uuid.UUID,
    current_user: CurrentUser,
    db: Annotated[AsyncSession, Depends(get_db)],
) -> DocumentStatusResponse:
    """Get the processing status of a document."""
    doc_repo = DocumentRepository(db)
    document = await doc_repo.get_by_id(document_id)
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )
    
    if document.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not authorized to access this document",
        )
    
    return DocumentStatusResponse(
        document_id=document_id,
        status=document.status,
        message=f"Document is {document.status}",
    )